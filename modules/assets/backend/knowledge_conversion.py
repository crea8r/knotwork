"""
File-to-Markdown conversion for Handbook uploads.

Supports: .md, .txt, .html, .csv, .pdf, .docx, .doc, .jpg, .png, .gif, .webp
Goal: best-effort structure preservation.  Expert reviews before saving.
"""
from __future__ import annotations

import csv
import io
import re
from pathlib import Path

IMAGE_EXTS = {'.jpg', '.jpeg', '.png', '.gif', '.webp'}
VIDEO_EXTS = {'.mp4', '.mov', '.avi', '.mkv', '.webm', '.flv', '.wmv', '.m4v', '.3gp'}
SUPPORTED_TEXT_EXTS = {'.md', '.txt', '.html', '.htm', '.csv', '.pdf', '.docx', '.doc'}
SUPPORTED_EXTS = SUPPORTED_TEXT_EXTS | IMAGE_EXTS  # VIDEO_EXTS not included


def convert_to_markdown(filename: str, content_bytes: bytes) -> tuple[str, str]:
    """Convert file bytes to Markdown.

    Returns (markdown_text, detected_format).
    Raises ValueError("video_not_supported") for video files.
    """
    suffix = Path(filename).suffix.lower()

    if suffix in VIDEO_EXTS:
        raise ValueError("video_not_supported")
    if suffix == ".md":
        return content_bytes.decode("utf-8", errors="replace"), "md"
    if suffix == ".txt":
        return _txt_to_md(filename, content_bytes), "txt"
    if suffix in (".html", ".htm"):
        return _html_to_md(filename, content_bytes), "html"
    if suffix == ".csv":
        return _csv_to_md(filename, content_bytes), "csv"
    if suffix == ".pdf":
        return _pdf_to_md(filename, content_bytes), "pdf"
    if suffix == ".docx":
        return _docx_to_md(filename, content_bytes), "docx"
    if suffix == ".doc":
        return _doc_to_md(filename, content_bytes), "doc"

    # Unknown — treat as plain text
    return content_bytes.decode("utf-8", errors="replace"), "txt"


def _stem(filename: str) -> str:
    """Return filename without extension, title-cased."""
    return Path(filename).stem.replace("-", " ").replace("_", " ").title()


def _txt_to_md(filename: str, content: bytes) -> str:
    text = content.decode("utf-8", errors="replace")
    heading = f"# {_stem(filename)}\n\n"
    return heading + text


def _html_to_md(filename: str, content: bytes) -> str:
    text = content.decode("utf-8", errors="replace")
    try:
        import markdownify
        return markdownify.markdownify(text, heading_style="ATX", strip=["script", "style", "nav", "footer"])
    except ImportError:
        pass
    # Fallback: strip tags
    import re
    no_tags = re.sub(r"<[^>]+>", "", text)
    return f"# {_stem(filename)}\n\n{no_tags.strip()}"


def _csv_to_md(filename: str, content: bytes) -> str:
    text = content.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = list(reader)
    if not rows:
        return f"# {_stem(filename)}\n\n(empty)"

    def _row(cells: list[str]) -> str:
        return "| " + " | ".join(str(c).replace("|", "\\|") for c in cells) + " |"

    lines = [f"# {_stem(filename)}\n", _row(rows[0]), _row(["---"] * len(rows[0]))]
    for row in rows[1:]:
        lines.append(_row(row))
    return "\n".join(lines)


def _pdf_to_md(filename: str, content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return f"# {_stem(filename)}\n\n(pypdf not installed — cannot extract PDF text)"

    reader = PdfReader(io.BytesIO(content))
    parts = [f"# {_stem(filename)}\n"]
    for page in reader.pages:
        page_text = page.extract_text() or ""
        for line in page_text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            # Heuristic: short ALL-CAPS lines → heading
            if stripped.isupper() and len(stripped) < 80:
                parts.append(f"\n## {stripped.title()}\n")
            else:
                parts.append(stripped)
    return "\n".join(parts)


def _docx_to_md(filename: str, content: bytes) -> str:
    try:
        import docx
    except ImportError:
        return f"# {_stem(filename)}\n\n(python-docx not installed — cannot extract .docx text)"

    doc = docx.Document(io.BytesIO(content))
    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            parts.append("")
            continue
        style = para.style.name if para.style else ""
        if "Heading 1" in style:
            parts.append(f"# {text}")
        elif "Heading 2" in style:
            parts.append(f"## {text}")
        elif "Heading 3" in style:
            parts.append(f"### {text}")
        else:
            parts.append(text)

    for table in doc.tables:
        if not table.rows:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        parts.append("\n| " + " | ".join(headers) + " |")
        parts.append("| " + " | ".join(["---"] * len(headers)) + " |")
        for row in table.rows[1:]:
            cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
            parts.append("| " + " | ".join(cells) + " |")
        parts.append("")

    return "\n".join(parts)


def _doc_to_md(filename: str, content: bytes) -> str:
    """Best-effort extraction of Word 97-2003 (.doc) OLE2 binary files."""
    warning = "> ⚠ Binary .doc format — save as .docx for better results.\n\n"
    heading = f"# {_stem(filename)}\n\n"

    # Attempt 1: maybe it's actually an OOXML file mislabeled as .doc
    try:
        import docx
        doc = docx.Document(io.BytesIO(content))
        parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        if parts:
            return warning + heading + "\n".join(parts)
    except Exception:
        pass

    # Attempt 2: OLE2 WordDocument stream via olefile
    try:
        import olefile
        if olefile.isOleFile(io.BytesIO(content)):
            ole = olefile.OleFileIO(io.BytesIO(content))
            if ole.exists("WordDocument"):
                raw = ole.openstream("WordDocument").read()
                # Word stores text in UTF-16LE; extract printable lines
                try:
                    decoded = raw.decode("utf-16-le", errors="ignore")
                    lines = [ln.strip() for ln in decoded.splitlines() if ln.strip() and ln.strip().isprintable()]
                    readable = [ln for ln in lines if len(ln) > 3]
                    if readable:
                        return warning + heading + "\n".join(readable[:200])
                except Exception:
                    pass
            ole.close()
    except Exception:
        pass

    # Attempt 3: ASCII run extraction (last resort)
    runs = re.findall(rb'[ -~]{8,}', content)
    lines = [r.decode("ascii", errors="ignore").strip() for r in runs]
    readable = [ln for ln in lines if ln and not ln.startswith("bjbj")]
    if readable:
        return warning + heading + "\n".join(readable[:200])

    return warning + heading + "(Could not extract text from this .doc file.)"


def suggested_path(filename: str, folder: str = "") -> str:
    """Return a suggested handbook path for an uploaded file."""
    stem = Path(filename).stem
    safe = stem.lower().replace(" ", "-").replace("_", "-")
    suffix = Path(filename).suffix.lower()
    extension = ".pptx" if suffix in {".ppt", ".pptx"} else ".md"
    if folder:
        return f"{folder.rstrip('/')}/{safe}{extension}"
    return f"{safe}{extension}"
